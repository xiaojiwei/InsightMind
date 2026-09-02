from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path('/Users/xiaojiwei/Downloads/九天-AlphaData-产品说明文档-浏览器复制.md')
OUTPUT = Path('/Users/xiaojiwei/Downloads/九天-AlphaData-产品说明文档-浏览器复制.docx')


def set_font(run, name='GB18030 Bitmap', size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run._element.rPr.rFonts.set(qn('w:cs'), name)
    run._element.rPr.rFonts.set(qn('w:hint'), 'eastAsia')
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = 'GB18030 Bitmap'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'GB18030 Bitmap')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h_tokens = {
        'Heading 1': (16, (46, 116, 181), 18, 10),
        'Heading 2': (13, (46, 116, 181), 14, 7),
        'Heading 3': (12, (31, 77, 120), 10, 5),
    }
    for style_name, (size, color, before, after) in h_tokens.items():
        style = doc.styles[style_name]
        style.font.name = 'GB18030 Bitmap'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'GB18030 Bitmap')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(header, after=0, line=1.0)
    set_font(header.add_run('AlphaData 产品说明文档'), size=9, color=(105, 112, 120))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(footer, after=0, line=1.0)
    set_font(footer.add_run('第 '), size=9, color=(105, 112, 120))
    add_page_field(footer)
    set_font(footer.add_run(' 页'), size=9, color=(105, 112, 120))

    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    first_title = True
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith('> '):
            p = doc.add_paragraph()
            set_spacing(p, after=3, line=1.15)
            set_font(p.add_run(line[2:].strip()), size=9.5, color=(100, 100, 100), italic=True)
            continue
        heading = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1 and first_title:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_spacing(p, before=12, after=12, line=1.0)
                set_font(p.add_run(text), size=22, color=(11, 37, 69), bold=True)
                first_title = False
            else:
                p = doc.add_paragraph(text, style=f'Heading {min(level, 3)}')
                set_spacing(p, before=h_tokens[f'Heading {min(level, 3)}'][2], after=h_tokens[f'Heading {min(level, 3)}'][3], line=1.15)
            continue
        bullet = re.match(r'^(\s*)-\s+(.+)$', line)
        number = re.match(r'^(\s*)\d+[.)]\s+(.+)$', line)
        if bullet or number:
            match = bullet or number
            indent = len(match.group(1).replace('\t', '    '))
            level = min(indent // 3, 2)
            style = ('List Bullet' if bullet else 'List Number') + (f' {level + 2}' if level else '')
            p = doc.add_paragraph(style=style)
            set_spacing(p, after=4, line=1.25)
            set_font(p.add_run(match.group(2).strip()), size=11)
            continue
        p = doc.add_paragraph()
        set_spacing(p, after=6, line=1.25)
        set_font(p.add_run(line.strip()), size=11)

    properties = doc.core_properties
    properties.title = '九天 AlphaData 产品说明文档'
    properties.subject = '浏览器复制保存的本地文档'
    properties.author = 'InsightMind'
    properties.comments = '由本地 Markdown 转换为 DOCX；飞书图片和视频以文本占位保留。'
    doc.save(OUTPUT)


if __name__ == '__main__':
    build_document()
